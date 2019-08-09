"""
Requirements:
lxml==4.4.0
Pillow==6.1.0
python-docx==0.8.10
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from tkinter import *
from tkinter import filedialog
import PIL.Image, PIL.ImageTk
import datetime


def inches_converter(pixels: int) -> float:
    """
    Converte pixels em polegadas, retornando o valor em polegadas.
    """

    inches = pixels/96
    return inches


def acquire_images() -> list:
    """
    Armazena todas as imagens do relatório em uma variável.
    """

    images = list(filedialog.askopenfilenames())
    images.sort(key=os.path.getmtime)  # Ordena os arquivos pela data da última modificação
    return images 


def docx_creator():
    """
    Cria o arquivo docx.
    """

    document = Document()

    actual_date = datetime.date.today()
    day, month, year = actual_date.day, actual_date.month, actual_date.year

    document.add_heading(f'Relatório do dia {day}/{month}/{year}', 0)
    
    images = acquire_images()

    for image in images:
        table = document.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Coluna 1'
        hdr_cells[1].text = 'Coluna 2'
        hdr_cells[2].text = 'Coluna 3'
        hdr_cells[3].text = 'Coluna 4'

        document.add_paragraph('Imagem', style='Caption')
        document.add_picture(image, width=Inches(inches_converter(360)))
        last_p = document.paragraphs[-1]
        last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_page_break()

    document.save(f'{filename.get()}.docx')
    os.startfile(f'{filename.get()}.docx')


root = Tk()
root.title('Gerenciador de relatórios')
root.geometry('700x650')

window_title = Label(root, text='Python Reports Maker', font=('Helvetica', 20), pady=20)
window_title.pack()

img = PIL.ImageTk.PhotoImage(PIL.Image.open('reports.png'))
panel = Label(root, image=img)
panel.pack()

filename_label = Label(root, text='Digite o nome do relatório a ser gerado', font=('Helvetica', 16))
filename_label.pack()

filename = Entry(root, text="Nome do relatório a ser gerado", width=30)
filename.pack()

button_label = Label(root, text='Selecione as imagens para gerar o relatório', 
                    font=('Helvetica', 16))
button_label.pack()

button = Button(root, text='Browse', font=('Helvetica', 14), command=docx_creator, pady=20, padx=20)
button.pack()

root.mainloop()
